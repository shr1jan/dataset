import os
import re
import hashlib
import queue
import threading
import time
import pdfplumber
import tkinter as tk
from tkinter import filedialog, messagebox, Listbox, Scrollbar, Frame, END, DISABLED, NORMAL
from tkinter import ttk
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import logging
from datetime import datetime
import json
import shutil
import warnings
import traceback
import sys

# Suppress pdfplumber warnings but keep critical ones
warnings.filterwarnings("ignore", category=UserWarning, message="CropBox missing from /Page, defaulting to MediaBox")

# Global variables
selected_pdf_paths = []
PRESERVE_AMENDMENTS = True
FORMAT_DATES = True
conversion_queue = queue.Queue()
abort_processing = False
processing_thread = None
backup_dir = "backup_documents"

def update_file_status(pdf_path, color):
    """Update a file's status color in the listbox"""
    filename = os.path.basename(pdf_path)
    items = file_listbox.get(0, tk.END)
    try:
        if filename in items:
            index = items.index(filename)
            file_listbox.itemconfig(index, {'fg': color})
        else:
            logging.warning(f"File {filename} not found in listbox")
    except Exception as e:
        logging.error(f"Error updating status for {filename}: {str(e)}")

# Set up robust logging
log_dir = "logs"
os.makedirs(log_dir, exist_ok=True)
os.makedirs(backup_dir, exist_ok=True)

log_file = os.path.join(log_dir, f"pdf_conversion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")
logging.basicConfig(
    filename=log_file,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Add console handler for critical errors
console_handler = logging.StreamHandler(sys.stdout)
console_handler.setLevel(logging.ERROR)
logging.getLogger().addHandler(console_handler)

def log_error(message, exception=None):
    """Log error with full stack trace and console output"""
    if exception:
        error_details = f"{message}: {str(exception)}\n{traceback.format_exc()}"
        logging.error(error_details)
        print(f"ERROR: {message}: {str(exception)}")
    else:
        logging.error(message)
        print(f"ERROR: {message}")

def create_document_backup(file_path):
    """Create backup of original document before processing"""
    if not os.path.exists(file_path):
        return None
        
    try:
        backup_file = os.path.join(backup_dir, os.path.basename(file_path))
        shutil.copy2(file_path, backup_file)
        logging.info(f"Created backup of {file_path} at {backup_file}")
        return backup_file
    except Exception as e:
        log_error(f"Failed to create backup for {file_path}", e)
        return None

def calculate_file_hash(file_path):
    """Calculate SHA-256 hash of file for integrity verification"""
    try:
        with open(file_path, 'rb') as f:
            file_hash = hashlib.sha256(f.read()).hexdigest()
        return file_hash
    except Exception as e:
        log_error(f"Failed to calculate hash for {file_path}", e)
        return None

def validate_pdf(pdf_path):
    """Validate that PDF file is readable and contains text"""
    try:
        logging.info(f"Validating PDF: {pdf_path}")
        with pdfplumber.open(pdf_path) as pdf:
            if len(pdf.pages) == 0:
                log_error(f"PDF has no pages: {pdf_path}")
                return False, "PDF has no pages"
                
            # Check at least the first, middle and last page for content
            pages_to_check = [0, len(pdf.pages)//2, -1]
            for page_idx in pages_to_check:
                page = pdf.pages[page_idx]
                text = page.extract_text()
                if not text or len(text.strip()) < 10:  # Arbitrary minimum text length
                    log_error(f"PDF page {page_idx+1} has insufficient text content: {pdf_path}")
                    return False, f"Page {page_idx+1} has insufficient text content"
                    
        return True, "PDF validated successfully"
    except Exception as e:
        log_error(f"PDF validation failed for {pdf_path}", e)
        return False, f"PDF validation error: {str(e)}"

def classify_line(line):
    """Classify line based on content patterns for styling"""
    line = line.strip()
    # Check for Notes: pattern first
    if re.match(r'^Notes\s*:', line, re.I):
        return "Normal"  # Treat Notes: as normal text
    # Check for special section formats with symbols - EXPANDED PATTERN
    elif re.match(r'^[♦◉]\s*\d+[A-Za-z]?\.', line):
        return "Heading 3"
    # Special case for list items with symbols (not section headers)
    elif re.match(r'^[♦◉]\s*\(\d+\)', line):
        return "Normal"  # Treat as normal text or list item
    elif re.match(r'^Schedule\b.*', line, re.I):
        return "Heading 5"
    elif re.match(r'^NEPAL.*ACT.*\d{4}', line, re.I):
        return "Title"
    elif re.match(r'^Date of (Authentication|Publication|Authentication and Publication|Royal Seal and Publication)\b.*', line, re.I):
        return "Subtitle"
    elif re.match(r'^AN ACT MADE TO.*', line, re.I):
        return "Subtitle"
    elif re.match(r'^Amendments\s*:?', line, re.I):
        return "Subtitle"
    elif re.match(r'^Preamble\s*:?', line, re.I):
        return "Heading 1"
    elif re.match(r'^Chapter\s*[-–]?\s*\d+', line, re.I):
        return "Heading 2"
    # New section formats
    elif re.match(r'^[♦◉]\s*\(\d+\)D?', line):
        return "Heading 3"
    elif re.match(r'^\d+[A-Za-z]?\.\s+', line):
        return "Heading 3"
    elif re.match(r'^\(\d+\)', line):
        return "Heading 4"
    elif re.match(r'^\d{4}\.\d{1,2}\.\d{1,2}.*', line):
         return "Normal"
    else:
        return "Normal"

def add_styled_paragraph(doc, text, style_tag, is_under_h5=False):
    """Add a styled paragraph to the document"""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style_tag]
    except KeyError:
        logging.warning(f"Style '{style_tag}' not found, using 'Normal' instead")
        p.style = doc.styles['Normal']

    run = p.add_run(text)

    if style_tag == "Heading 5":
        pass
    elif style_tag == "Heading 4":
        p.paragraph_format.left_indent = Inches(0.3)
    elif style_tag == "Normal" and is_under_h5:
        p.paragraph_format.left_indent = Inches(0.3)
    elif style_tag == "Normal":
        p.paragraph_format.left_indent = Inches(0)
    elif style_tag == "Title":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.font.size = Pt(16)
        run.font.bold = True
    elif style_tag == "Subtitle":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run.font.size = Pt(14)
        run.font.italic = True

    return p

def add_table_to_doc(doc, table_data):
    """Add a table to the document with the given data."""
    if not table_data or not table_data[0]:
        return None
    
    num_rows = len(table_data)
    num_cols = max(len(row) for row in table_data)
    
    # Ensure table data is valid
    if num_rows == 0 or num_cols == 0:
        logging.warning("Attempted to add empty table - skipping")
        return None
    
    try:
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Fill the table with data
        for i, row in enumerate(table_data):
            for j, cell_text in enumerate(row):
                if j < num_cols:  # Ensure we don't exceed the number of columns
                    table.cell(i, j).text = cell_text
        
        return table
    except Exception as e:
        log_error("Failed to add table to document", e)
        return None

def is_likely_table_row(line):
    """Check if a line is likely part of a table based on patterns."""
    # Check for multiple cell separators like | or tab characters
    if '|' in line and line.count('|') >= 2:
        return True
    if '\t' in line and line.count('\t') >= 1:
        return True
    # Check for patterns like "Column1   Column2   Column3"
    if re.search(r'\S+\s{2,}\S+\s{2,}\S+', line):
        return True
    return False

def extract_table_from_lines(lines, start_idx):
    """Extract table data from consecutive lines that appear to be a table."""
    table_data = []
    i = start_idx
    
    while i < len(lines) and is_likely_table_row(lines[i]):
        row = lines[i]
        # Split by pipe if present
        if '|' in row:
            cells = [cell.strip() for cell in row.split('|')]
            # Remove empty cells at the beginning and end if they're just artifacts of the pipe splitting
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
        # Split by tabs if present
        elif '\t' in row:
            cells = [cell.strip() for cell in row.split('\t')]
        # Split by multiple spaces
        else:
            cells = re.split(r'\s{2,}', row.strip())
        
        table_data.append(cells)
        i += 1
    
    return table_data, i - start_idx

def convert_pdf_to_docx(pdf_path, output_dir=None, progress_callback=None):
    """Convert PDF to structured DOCX with progress updates and validation"""
    doc = Document()
    url_pattern = re.compile(r'(?:https?://|www\.)\S+')
    translation_pattern = re.compile(r'\s*\((Official|Unofficial)\s+Translation\)\s*', re.I)
    is_within_heading_5 = False
    last_p = None
    last_tag = None
    is_first_line_of_document = True
    is_within_amendments = False
    is_within_subsection = False  # Track if we're inside a subsection like (a), (b), etc.
    is_processing_table = False
    table_buffer = []
    all_tables = []
    document_stats = {
        "pages_processed": 0,
        "sections_found": 0,
        "tables_found": 0,
        "total_pages": 0,
        "headings": {
            "title": 0,
            "subtitle": 0,
            "h1": 0,
            "h2": 0,
            "h3": 0,
            "h4": 0, 
            "h5": 0
        }
    }
    
    # Create backup of the original document
    create_document_backup(pdf_path)
    
    # Validate PDF before processing
    is_valid, validation_message = validate_pdf(pdf_path)
    if not is_valid:
        return None, validation_message, document_stats
    
    try:
        logging.info(f"Starting conversion of: {os.path.basename(pdf_path)}")
        
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            document_stats["total_pages"] = total_pages
            logging.info(f"PDF has {total_pages} pages")
            
            # Process each page for text
            for page_num, page in enumerate(pdf.pages, 1):
                if abort_processing:
                    logging.warning(f"Processing aborted for {pdf_path}")
                    return None, "Processing aborted by user", document_stats
                
                # Update progress
                if progress_callback:
                    progress_callback(page_num, total_pages, f"Processing page {page_num}/{total_pages}")
                
                text = page.extract_text()
                if not text:
                    logging.warning(f"No text found on page {page_num}")
                    continue
                
                document_stats["pages_processed"] += 1
                lines = text.split("\n")
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    original_line = line
                    
                    # Remove translation markers
                    line = translation_pattern.sub('', line)
                    
                    if not line:
                        if is_within_heading_5 and not is_within_amendments:
                            current_p = add_styled_paragraph(doc, "", "Normal", is_under_h5=True)
                            last_p = current_p
                            last_tag = "Normal"
                        i += 1
                        continue
                    
                    if re.fullmatch(r'\d+', line):
                        i += 1
                        continue
                    
                    # Check if this line is likely the start of a table
                    if is_likely_table_row(line) and not is_processing_table:
                        is_processing_table = True
                        table_data, rows_consumed = extract_table_from_lines(lines, i)
                        if table_data and len(table_data) > 1:  # Ensure it's actually a table with multiple rows
                            add_table_to_doc(doc, table_data)
                            all_tables.append(table_data)  # Store table data for later reference
                            document_stats["tables_found"] += 1
                            i += rows_consumed
                            is_processing_table = False
                            continue
                        else:
                            # Not a real table, process as normal text
                            is_processing_table = False
                    
                    # Process as normal text if not a table
                    line_after_url_removal = url_pattern.sub('', line).strip()
                    line_after_url_removal = translation_pattern.sub('', line_after_url_removal)
                    
                    if not line_after_url_removal:
                        i += 1
                        continue
                    
                    if is_first_line_of_document:
                        tag = "Title"
                        current_p = add_styled_paragraph(doc, line_after_url_removal, tag)
                        document_stats["headings"]["title"] += 1
                        last_p = current_p
                        last_tag = tag
                        is_first_line_of_document = False
                        is_within_amendments = False
                        is_within_heading_5 = False
                        i += 1
                        continue

                    tag = classify_line(line_after_url_removal)

                    # Check if this is a subsection marker like (a), (b), etc.
                    is_letter_subsection = re.match(r'^\([a-z]\)', line_after_url_removal)
                    if is_letter_subsection:
                        is_within_subsection = True
                    
                    # Check if this is a numbered item inside a subsection
                    is_numbered_item_in_subsection = is_within_subsection and re.match(r'^\(\d+\)', line_after_url_removal)
                    
                    # If it's a numbered item inside a subsection, treat it as normal text
                    if is_numbered_item_in_subsection:
                        tag = "Normal"

                    if is_within_amendments:
                        if tag in ["Title", "Heading 1", "Heading 2"] or \
                           (tag == "Subtitle" and not re.match(r'^Amendments\s*:?', line_after_url_removal, re.I)):
                            is_within_amendments = False
                        else:
                            current_p = add_styled_paragraph(doc, line_after_url_removal, "Subtitle")
                            document_stats["headings"]["subtitle"] += 1
                            last_p = current_p
                            last_tag = "Subtitle"
                            i += 1
                            continue

                    is_date_line = re.match(r'^\d{4}\.\d{1,2}\.\d{1,2}', original_line)
                    is_prev_date_subtitle = last_p is not None and last_tag == "Subtitle" and \
                                            re.match(r'^Date of (Authentication|Publication|Authentication and Publication|Royal Seal and Publication)', last_p.text.split('\n')[0], re.I)
                    if is_prev_date_subtitle and is_date_line:
                        last_p.add_run(f"\n{original_line}")
                        i += 1
                        continue

                    current_p = None

                    if tag == "Subtitle" and re.match(r'^Amendments\s*:?', line_after_url_removal, re.I):
                        is_within_amendments = True
                        is_within_heading_5 = False
                        current_p = add_styled_paragraph(doc, line_after_url_removal, tag)
                        document_stats["headings"]["subtitle"] += 1
                    elif tag == "Heading 5":
                        is_within_heading_5 = True
                        current_p = add_styled_paragraph(doc, line_after_url_removal, tag)
                        document_stats["headings"]["h5"] += 1
                    elif tag in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
                        is_within_heading_5 = False
                        
                        # Reset subsection tracking when we hit a new section
                        if tag in ["Heading 1", "Heading 2", "Heading 3"]:
                            is_within_subsection = False
                            
                        if tag == "Heading 3":
                            document_stats["sections_found"] += 1
                            document_stats["headings"]["h3"] += 1
                            # Check for standard section format (number followed by dot)
                            sec_match = re.match(r'^(\d+[A-Za-z]?)\.\s*(.*)', line_after_url_removal)
                            # Check for special section formats with symbols - UPDATED PATTERN
                            symbol_sec_match = re.match(r'^([♦◉])\s*(\d+[A-Za-z]?)\.?\s*(.*)', line_after_url_removal)
                            # New pattern for list items with symbols
                            symbol_list_match = re.match(r'^([♦◉])\s*\((\d+)\)\s*(.*)', line_after_url_removal)
                            
                            if sec_match:
                                sec_num, sec_body = sec_match.groups()
                                sec_body = sec_body.strip()
                                parts = re.split(r'\s*(?=\(\d+\))', sec_body, maxsplit=1)
                                section_title = parts[0].strip()
                                current_p = add_styled_paragraph(doc, f"Section {sec_num}: {section_title}", "Heading 3")
                                if len(parts) > 1:
                                    first_subsection_text = parts[1].strip()
                                    sub_match = re.match(r'^\((\d+)\)\s*(.*)', first_subsection_text)
                                    if sub_match:
                                        sub_num, sub_title = sub_match.groups()
                                        # Check if this is a lettered subsection like (a), (b)
                                        if re.match(r'^[a-z]$', sub_num):
                                            is_within_subsection = True
                                            current_p = add_styled_paragraph(doc, f"({sub_num}) {sub_title.strip()}", "Heading 4")
                                            document_stats["headings"]["h4"] += 1
                                        else:
                                            # If we're inside a lettered subsection, treat numbered items as normal text
                                            if is_within_subsection:
                                                current_p = add_styled_paragraph(doc, f"({sub_num}) {sub_title.strip()}", "Normal")
                                            else:
                                                # Format long subsections with the number separated from content
                                                current_p = add_styled_paragraph(doc, f"Subsection ({sub_num}):", "Heading 4")
                                                document_stats["headings"]["h4"] += 1
                                                current_p = add_styled_paragraph(doc, f"{sub_title.strip()}", "Normal")
                                                tag = "Normal"
                                    else:
                                        current_p = add_styled_paragraph(doc, first_subsection_text, "Normal")
                                        tag = "Normal"
                            elif symbol_sec_match:
                                symbol, sec_num, sec_body = symbol_sec_match.groups()
                                sec_body = sec_body.strip()
                                section_format = f"{symbol}{sec_num}"
                                
                                parts = re.split(r'\s*(?=\(\d+\))', sec_body, maxsplit=1)
                                section_title = parts[0].strip()
                                
                                current_p = add_styled_paragraph(doc, f"Section {section_format}: {section_title}", "Heading 3")
                                
                                if len(parts) > 1:
                                    first_subsection_text = parts[1].strip()
                                    sub_match = re.match(r'^\((\d+)\)\s*(.*)', first_subsection_text)
                                    if sub_match:
                                        sub_num, sub_text = sub_match.groups()
                                        # Check if we're inside a lettered subsection
                                        if is_within_subsection:
                                            current_p = add_styled_paragraph(doc, f"({sub_num}) {sub_text.strip()}", "Normal")
                                            tag = "Normal"
                                        else:
                                            # Format long subsections with the number separated from content
                                            current_p = add_styled_paragraph(doc, f"Subsection ({sub_num}):", "Heading 4")
                                            document_stats["headings"]["h4"] += 1
                                            current_p = add_styled_paragraph(doc, f"{sub_text.strip()}", "Normal")
                                            tag = "Normal"
                                    else:
                                        current_p = add_styled_paragraph(doc, first_subsection_text, "Normal")
                                        tag = "Normal"
                            elif symbol_list_match:
                                # Handle list items with symbols
                                symbol, list_num, list_text = symbol_list_match.groups()
                                current_p = add_styled_paragraph(doc, f"{symbol} ({list_num}) {list_text.strip()}", "Normal")
                                tag = "Normal"
                            else:
                                current_p = add_styled_paragraph(doc, line_after_url_removal, "Heading 3")
                        elif tag == "Heading 2":
                            document_stats["headings"]["h2"] += 1
                            chap_match = re.match(r'^Chapter\s*[-–]?\s*(\d+)\s*(.*)', line_after_url_removal, re.I)
                            if chap_match:
                                chap_num, chap_title = chap_match.groups()
                                full_title = f"Chapter {chap_num.strip()}: {chap_title.strip()}"
                                current_p = add_styled_paragraph(doc, full_title, "Heading 2")
                            else:
                                current_p = add_styled_paragraph(doc, line_after_url_removal, "Heading 2")
                        elif tag == "Heading 1":
                            document_stats["headings"]["h1"] += 1
                            current_p = add_styled_paragraph(doc, line_after_url_removal, tag)
                        else:
                            if tag == "Title":
                                document_stats["headings"]["title"] += 1
                            elif tag == "Subtitle":
                                document_stats["headings"]["subtitle"] += 1
                            current_p = add_styled_paragraph(doc, line_after_url_removal, tag)
                    elif tag == "Normal":
                        current_p = add_styled_paragraph(doc, line_after_url_removal, "Normal", is_under_h5=is_within_heading_5)

                    if current_p:
                        last_p = current_p
                        last_tag = tag
                    
                    i += 1
            
            # Add any tables that were detected but not already processed
            if all_tables:
                doc.add_paragraph("").add_run("Tables from Document:").bold = True
                for table_data in all_tables:
                    # Filter out empty rows and cells
                    filtered_table = [[cell if cell else "" for cell in row] for row in table_data if any(cell for cell in row)]
                    if filtered_table:
                        add_table_to_doc(doc, filtered_table)
                        doc.add_paragraph()  # Add space after table

        # Add document metadata
        doc.core_properties.title = os.path.basename(os.path.splitext(pdf_path)[0])
        doc.core_properties.created = datetime.now()
        doc.core_properties.comments = f"Converted from PDF by Structured Document Converter v2.0"

        # Determine output path
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            output_filename = os.path.basename(os.path.splitext(pdf_path)[0]) + "_structured.docx"
            output_path = os.path.join(output_dir, output_filename)
        else:
            output_path = os.path.splitext(pdf_path)[0] + "_structured.docx"
            
        # Final progress update
        if progress_callback:
            progress_callback(total_pages, total_pages, "Saving document...")
            
        # Save document with error handling
        try:
            doc.save(output_path)
        except Exception as e:
            log_error(f"Failed to save document {output_path}", e)
            return None, f"Failed to save document: {str(e)}", document_stats
            
        # Verify output document
        if not os.path.exists(output_path) or os.path.getsize(output_path) < 1000:  # Arbitrary minimum size
            log_error(f"Output document verification failed for {output_path}: File too small or missing")
            return None, "Output document verification failed", document_stats
            
        # Log success
        logging.info(f"Successfully converted: {pdf_path} to {output_path}")
        logging.info(f"Document statistics: {json.dumps(document_stats)}")
        
        return output_path, "Success", document_stats
        
    except Exception as e:
        error_msg = f"Error processing: {pdf_path} - {str(e)}"
        log_error(error_msg, e)
        return None, error_msg, document_stats

def verify_docx_integrity(docx_path, stats):
    """Verify that the DOCX file has expected structure based on stats"""
    try:
        doc = Document(docx_path)
        
        # Check basic structure
        if len(doc.paragraphs) < 10:  # Arbitrary minimum
            return False, "Document has too few paragraphs"
            
        # Count headings to verify against stats
        heading_counts = {
            "Heading 1": 0,
            "Heading 2": 0,
            "Heading 3": 0,
            "Title": 0
        }
        
        for p in doc.paragraphs:
            if p.style.name in heading_counts:
                heading_counts[p.style.name] += 1
        
        # Compare with stats
        if stats["headings"]["h1"] > 0 and heading_counts["Heading 1"] == 0:
            return False, "Missing Heading 1 elements that were in the source"
            
        if stats["headings"]["h2"] > 0 and heading_counts["Heading 2"] == 0:
            return False, "Missing Heading 2 elements that were in the source"
            
        if stats["headings"]["h3"] > 0 and heading_counts["Heading 3"] == 0:
            return False, "Missing Heading 3 elements that were in the source"
            
        if stats["headings"]["title"] > 0 and heading_counts["Title"] == 0:
            return False, "Missing Title element that was in the source"
            
        # Check tables
        if stats["tables_found"] > 0 and len(doc.tables) == 0:
            return False, "Missing tables that were in the source"
            
        return True, "Document structure verified"
    except Exception as e:
        log_error(f"Document verification failed for {docx_path}", e)
        return False, f"Document verification error: {str(e)}"

def process_queue():
    """Process files from the queue with progress updates"""
    global abort_processing
    
    while not conversion_queue.empty() and not abort_processing:
        pdf_path, output_dir = conversion_queue.get()
        try:
            # Update UI to show current file
            root.after(0, lambda: update_status(f"Processing: {os.path.basename(pdf_path)}...", "blue"))
            root.after(0, lambda: progress_bar.configure(value=0))
            
            # Process the file with progress updates
            output_path, status_msg, doc_stats = convert_pdf_to_docx(
                pdf_path, 
                output_dir,
                progress_callback=lambda current, total, msg: root.after(0, 
                    lambda c=current, t=total, m=msg: update_progress(c, t, m))
            )
            
            if output_path:
                # Verify the document
                is_valid, verify_msg = verify_docx_integrity(output_path, doc_stats)
                if is_valid:
                    root.after(0, lambda p=pdf_path: file_listbox.itemconfig(
                        file_listbox.get(0, tk.END).index(os.path.basename(p)), 
                        {'fg': 'green'}
                    ))
                    logging.info(f"✓ {os.path.basename(pdf_path)} - Converted successfully and verified")
                    
                    # Generate and save conversion report
                    report_path = os.path.splitext(output_path)[0] + "_report.json"
                    with open(report_path, 'w') as f:
                        json.dump({
                            "source_file": pdf_path,
                            "output_file": output_path,
                            "conversion_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            "status": "success",
                            "document_statistics": doc_stats,
                            "verification": "passed"
                        }, f, indent=4)
                else:
                    root.after(0, lambda p=pdf_path: file_listbox.itemconfig(
                        file_listbox.get(0, tk.END).index(os.path.basename(p)), 
                        {'fg': 'orange'}
                    ))
                    logging.warning(f"⚠ {os.path.basename(pdf_path)} - Converted but verification failed: {verify_msg}")
                    
                    # If verification failed, notify user
                    root.after(0, lambda p=pdf_path, msg=verify_msg: messagebox.showwarning(
                        "Document Verification Warning", 
                        f"The document {os.path.basename(p)} was converted but failed verification: {msg}\n\nPlease review the output file manually."
                    ))
                    
                    # Save report even for warnings
                    report_path = os.path.splitext(output_path)[0] + "_report.json"
                    with open(report_path, 'w') as f:
                        json.dump({
                            "source_file": pdf_path,
                            "output_file": output_path,
                            "conversion_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            "status": "warning",
                            "warning_message": verify_msg,
                            "document_statistics": doc_stats,
                            "verification": "warning"
                        }, f, indent=4)
            else:
                root.after(0, lambda p=pdf_path: file_listbox.itemconfig(
                    file_listbox.get(0, tk.END).index(os.path.basename(p)), 
                    {'fg': 'red'}
                ))
                logging.error(f"❌ {os.path.basename(pdf_path)} - Conversion failed: {status_msg}")
                
                # Save failure report
                report_dir = os.path.dirname(pdf_path) if not output_dir else output_dir
                report_path = os.path.join(report_dir, os.path.splitext(os.path.basename(pdf_path))[0] + "_error_report.json")
                with open(report_path, 'w') as f:
                    json.dump({
                        "source_file": pdf_path,
                        "conversion_attempt_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        "status": "error",
                        "error_message": status_msg,
                        "partial_document_statistics": doc_stats
                    }, f, indent=4)
                
        except Exception as e:
            log_error(f"Failed to process {pdf_path}", e)
            root.after(0, lambda p=pdf_path: file_listbox.itemconfig(
                file_listbox.get(0, tk.END).index(os.path.basename(p)), 
                {'fg': 'red'}
            ))
        finally:
            conversion_queue.task_done()
    
    # Update UI when all files are processed
    root.after(0, processing_complete)

def update_progress(current, total, message=""):
    """Update progress bar and status"""
    if total > 0:
        percentage = int((current / total) * 100)
        progress_bar["value"] = percentage
        progress_label.config(text=f"{percentage}% - {message}")
    root.update_idletasks()

def update_status(message, color="black"):
    """Update status label with message and color"""
    status_label.config(text=message, fg=color)
    root.update_idletasks()

def start_processing():
    """Start processing files in a separate thread"""
    global processing_thread, abort_processing
    
    if not selected_pdf_paths:
        messagebox.showwarning("No Files", "Please select PDF files first.")
        return
    
    # Disable buttons during processing
    select_button.config(state=DISABLED)
    convert_button.config(state=DISABLED)
    abort_button.config(state=NORMAL)
    
    # Clear any previous abort flag
    abort_processing = False
    
    # Add files to the queue
    for path in selected_pdf_paths:
        conversion_queue.put((path, output_dir_var.get() if output_dir_var.get() else None))
        
    # Update UI
    update_status("Starting conversion...", "blue")
    progress_bar["value"] = 0
    
    # Start processing thread
    processing_thread = threading.Thread(target=process_queue, daemon=True)
    processing_thread.start()
    
    # Check thread status periodically
    check_thread_status()

def check_thread_status():
    """Check if processing thread is still running and update UI accordingly"""
    if processing_thread and processing_thread.is_alive():
        # Still processing, check again in 100ms
        root.after(100, check_thread_status)
    else:
        # Processing complete or was never started
        if not abort_processing and conversion_queue.empty():
            processing_complete()

def processing_complete():
    """Update UI when processing is complete"""
    select_button.config(state=NORMAL)
    convert_button.config(state=NORMAL)
    abort_button.config(state=DISABLED)
    update_status("Conversion complete", "green")
    progress_bar["value"] = 100
    progress_label.config(text="100% - Complete")

def abort_conversion():
    """Abort the current conversion process"""
    global abort_processing
    
    if messagebox.askyesno("Abort Conversion", "Are you sure you want to abort the current conversion process?"):
        abort_processing = True
        update_status("Aborting... Please wait", "red")
        
        # Wait for thread to finish current file
        if processing_thread and processing_thread.is_alive():
            processing_thread.join(timeout=1.0)  # Wait up to 1 second
            
        # Clear queue
        while not conversion_queue.empty():
            try:
                conversion_queue.get_nowait()
                conversion_queue.task_done()
            except queue.Empty:
                break
                
        # Update UI
        select_button.config(state=NORMAL)
        convert_button.config(state=NORMAL)
        abort_button.config(state=DISABLED)
        update_status("Conversion aborted", "red")

def select_files():
    """Select PDF files for conversion"""
    global selected_pdf_paths
    selected_pdf_paths.clear()
    file_listbox.delete(0, tk.END)
    status_label.config(text="")

    file_paths = filedialog.askopenfilenames(
        title="Select PDF Files", 
        filetypes=[("PDF Files", "*.pdf")]
    )

    if not file_paths:
        convert_button.config(state=DISABLED)
        return

    # Validate files before adding to list
    for path in file_paths:
        # Basic file checks
        if not os.path.exists(path):
            messagebox.showwarning("File Not Found", f"The file {os.path.basename(path)} does not exist.")
            continue
            
        if os.path.getsize(path) == 0:
            messagebox.showwarning("Empty File", f"The file {os.path.basename(path)} is empty.")
            continue
            
        # Try to open with pdfplumber to validate
        try:
            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)
                if page_count == 0:
                    messagebox.showwarning("Invalid PDF", f"The file {os.path.basename(path)} has no pages.")
                    continue
        except Exception as e:
            messagebox.showwarning("Invalid PDF", f"The file {os.path.basename(path)} could not be opened: {str(e)}")
            continue
            
        # If all checks pass, add to list
        selected_pdf_paths.append(path)
        file_listbox.insert(tk.END, os.path.basename(path))

    if selected_pdf_paths:
        status_label.config(text=f"{len(selected_pdf_paths)} PDF(s) selected and validated", fg="green")
        convert_button.config(state=NORMAL)
    else:
        status_label.config(text="No valid PDF files selected", fg="red")
        convert_button.config(state=DISABLED)

def select_output_dir():
    """Select output directory for converted files"""
    output_dir = filedialog.askdirectory(title="Select Output Directory")
    if output_dir:
        output_dir_var.set(output_dir)
        output_dir_label.config(text=f"Output: {output_dir}")

def create_config():
    """Create default configuration file if it doesn't exist"""
    config_path = "pdf_converter_config.json"
    if not os.path.exists(config_path):
        default_config = {
            "preserve_amendments": True,
            "format_dates": True,
            "default_output_dir": "",
            "log_level": "INFO",
            "max_threads": 1,
            "auto_verify": True,
            "backup_files": True
        }
        with open(config_path, 'w') as f:
            json.dump(default_config, f, indent=4)
        return default_config
    else:
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except Exception as e:
            log_error(f"Failed to load configuration file", e)
            return {
                "preserve_amendments": True,
                "format_dates": True,
                "default_output_dir": "",
                "log_level": "INFO",
                "max_threads": 1,
                "auto_verify": True,
                "backup_files": True
            }

def show_about():
    """Show about dialog with version and credits"""
    messagebox.showinfo(
        "About PDF to Structured DOCX Converter",
        "PDF to Structured DOCX Converter v2.0\n\n"
        "This application converts PDF documents to structured DOCX format "
        "with proper heading hierarchy and formatting.\n\n"
        "© 2025 All Rights Reserved\n\n"
        "Log file: " + os.path.abspath(log_file)
    )

def show_help():
    """Show help dialog with usage instructions"""
    messagebox.showinfo(
        "Help - PDF to Structured DOCX Converter",
        "How to use this application:\n\n"
        "1. Click 'Select PDF Files' to choose one or more PDF files for conversion.\n"
        "2. Optionally select an output directory (defaults to same location as PDF).\n"
        "3. Click 'Convert to DOCX' to start the conversion process.\n"
        "4. Monitor progress in the status area below.\n"
        "5. Green entries indicate successful conversion.\n"
        "6. Orange entries indicate successful conversion with verification warnings.\n"
        "7. Red entries indicate failed conversion.\n\n"
        "For each converted file, a report file is generated with details of the conversion process."
    )

def show_log():
    """Open log file in default text editor"""
    try:
        if os.path.exists(log_file):
            os.startfile(log_file) if sys.platform == 'win32' else os.system(f'open "{log_file}"')
        else:
            messagebox.showinfo("Log File", "No log file exists yet.")
    except Exception as e:
        messagebox.showerror("Error", f"Could not open log file: {str(e)}")

def on_closing():
    """Handle window closing event"""
    global abort_processing
    
    if processing_thread and processing_thread.is_alive():
        if messagebox.askyesno("Quit", "Conversion is in progress. Are you sure you want to quit?"):
            abort_processing = True
            root.destroy()
    else:
        root.destroy()

# Create the main window with improved design
root = tk.Tk()
root.title("PDF to Structured DOCX Converter v2.0")
root.geometry("800x600")
root.minsize(700, 500)
root.protocol("WM_DELETE_WINDOW", on_closing)

# Load configuration
config = create_config()

# Create main frames
top_frame = Frame(root, padx=10, pady=10)
top_frame.pack(fill=tk.X)

file_frame = Frame(root, padx=10, pady=5)
file_frame.pack(fill=tk.BOTH, expand=True)

status_frame = Frame(root, padx=10, pady=10)
status_frame.pack(fill=tk.X, side=tk.BOTTOM)

# Create menu bar
menu_bar = tk.Menu(root)
root.config(menu=menu_bar)

file_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="Select PDF Files", command=select_files)
file_menu.add_command(label="Select Output Directory", command=select_output_dir)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=on_closing)

tools_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Tools", menu=tools_menu)
tools_menu.add_command(label="View Log File", command=show_log)

help_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Help", menu=help_menu)
help_menu.add_command(label="User Guide", command=show_help)
help_menu.add_command(label="About", command=show_about)

# Create buttons in top frame
select_button = ttk.Button(top_frame, text="Select PDF Files", command=select_files)
select_button.pack(side=tk.LEFT, padx=5)

convert_button = ttk.Button(top_frame, text="Convert to DOCX", command=start_processing, state=DISABLED)
convert_button.pack(side=tk.LEFT, padx=5)

abort_button = ttk.Button(top_frame, text="Abort Conversion", command=abort_conversion, state=DISABLED)
abort_button.pack(side=tk.LEFT, padx=5)

output_dir_var = tk.StringVar(value=config.get("default_output_dir", ""))
output_button = ttk.Button(top_frame, text="Select Output Directory", command=select_output_dir)
output_button.pack(side=tk.LEFT, padx=5)

output_dir_label = ttk.Label(top_frame, text="Output: Default (same as PDF)")
if output_dir_var.get():
    output_dir_label.config(text=f"Output: {output_dir_var.get()}")
output_dir_label.pack(side=tk.LEFT, padx=5)

# Create file listbox with scrollbar
file_listbox_label = ttk.Label(file_frame, text="Files to Process:")
file_listbox_label.pack(anchor=tk.W)

file_list_frame = Frame(file_frame)
file_list_frame.pack(fill=tk.BOTH, expand=True)

scrollbar = ttk.Scrollbar(file_list_frame)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

file_listbox = tk.Listbox(file_list_frame, yscrollcommand=scrollbar.set, selectmode=tk.EXTENDED, font=("Courier", 10))
file_listbox.pack(fill=tk.BOTH, expand=True)
scrollbar.config(command=file_listbox.yview)

# Create status indicators in status frame
status_label = tk.Label(status_frame, text="", font=("Arial", 10))
status_label.pack(anchor=tk.W)

progress_frame = Frame(status_frame)
progress_frame.pack(fill=tk.X, pady=5)

progress_bar = ttk.Progressbar(progress_frame, orient=tk.HORIZONTAL, length=100, mode='determinate')
progress_bar.pack(fill=tk.X, side=tk.LEFT, expand=True)

progress_label = ttk.Label(progress_frame, text="0%")
progress_label.pack(side=tk.RIGHT, padx=5)

# Set up error handling for the entire application
def show_error(exception_type, exception_value, exception_traceback):
    """Global exception handler"""
    error_msg = f"An unhandled error occurred:\n{exception_type.__name__}: {exception_value}"
    log_error("Unhandled exception", exception_value)
    messagebox.showerror("Critical Error", error_msg)
    
    # Write full traceback to log
    import traceback
    logging.error("Full traceback:")
    logging.error(''.join(traceback.format_tb(exception_traceback)))

# Set up global exception handler
sys.excepthook = show_error

# Start the main loop with better exception handling
if __name__ == "__main__":
    try:
        root.mainloop()
    except Exception as e:
        log_error("Fatal error in main loop", e)